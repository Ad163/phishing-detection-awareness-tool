"""
Converts a dissertation chapter written in Markdown into a .docx formatted
to the PROM02 module's submission requirements (Info.docx):
    - A4 paper
    - Clear 12-point font (Times New Roman)
    - At least 20mm margin all round
    - One-and-a-half line spacing throughout body text
    - Consecutively numbered pages

Usage:
    python md_to_docx.py <input.md> <output.docx> [--chapter-start-page 1]

Markdown conventions expected:
    # Title            -> Heading 1 (chapter title)
    ## 1.1 Section      -> Heading 2
    ### 1.1.1 Subsection -> Heading 3
    **bold text**        -> bold run
    blank-line separated -> new paragraph
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(12)
LINE_SPACING = 1.5
MARGIN = Mm(25)  # 25mm > the 20mm minimum required, gives a safety margin


def add_page_number_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def set_base_style(doc):
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE
    style.paragraph_format.line_spacing = LINE_SPACING
    style.paragraph_format.space_after = Pt(8)
    # ensure east-asian font attr doesn't override for special characters
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_NAME)

    for heading_style, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        hstyle = doc.styles[heading_style]
        hstyle.font.name = FONT_NAME
        hstyle.font.size = Pt(size)
        hstyle.font.bold = True
        hstyle.font.color.rgb = None
        hstyle.paragraph_format.line_spacing = LINE_SPACING
        hstyle.paragraph_format.space_before = Pt(18)
        hstyle.paragraph_format.space_after = Pt(10)


def set_page_setup(section):
    section.page_height = Mm(297)  # A4
    section.page_width = Mm(210)
    section.left_margin = MARGIN
    section.right_margin = MARGIN
    section.top_margin = MARGIN
    section.bottom_margin = MARGIN


def add_bold_runs(paragraph, text):
    """Split on **bold** markers and add runs accordingly."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def is_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.endswith("|") and s.count("|") >= 2


def is_separator_row(line: str) -> bool:
    s = line.strip().replace("|", "")
    return bool(s) and all(c in "-: " for c in s)


def parse_table_row(line: str):
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [cell.strip() for cell in s.split("|")]


def add_table(doc, rows):
    """rows: list of lists of cell strings; first row is the header."""
    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"
    for r, row_cells in enumerate(rows):
        for c in range(n_cols):
            cell_text = row_cells[c] if c < len(row_cells) else ""
            cell = table.cell(r, c)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(cell_text)
            run.font.name = FONT_NAME
            run.font.size = Pt(11)
            if r == 0:
                run.bold = True
    # small space after the table
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def convert(md_path: Path, docx_path: Path):
    text = md_path.read_text(encoding="utf-8")
    lines = text.split("\n")

    doc = Document()
    set_base_style(doc)
    set_page_setup(doc.sections[0])

    # footer with page number, centred
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number_field(footer_para)

    buffer = []

    def flush_paragraph():
        if buffer:
            joined = " ".join(buffer).strip()
            if joined:
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing = LINE_SPACING
                p.paragraph_format.space_after = Pt(8)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                add_bold_runs(p, joined)
            buffer.clear()

    i = 0
    n = len(lines)
    while i < n:
        raw_line = lines[i]
        line = raw_line.rstrip()

        if is_table_row(line) and i + 1 < n and is_separator_row(lines[i + 1]):
            flush_paragraph()
            table_rows = [parse_table_row(line)]
            i += 2  # skip header + separator
            while i < n and is_table_row(lines[i]):
                table_rows.append(parse_table_row(lines[i]))
                i += 1
            add_table(doc, table_rows)
            continue

        if line.startswith("### "):
            flush_paragraph()
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            flush_paragraph()
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            flush_paragraph()
            doc.add_heading(line[2:], level=1)
        elif line.strip() == "":
            flush_paragraph()
        else:
            buffer.append(line.strip())
        i += 1
    flush_paragraph()

    doc.save(docx_path)
    print(f"Saved {docx_path}")


if __name__ == "__main__":
    md_file = Path(sys.argv[1])
    out_file = Path(sys.argv[2])
    convert(md_file, out_file)
